import io
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate
from fastapi import HTTPException

from models import MeetingModel, Section, SectionDates, TemplateSpec


BASE_DIR = Path(__file__).resolve().parent


def render_docx(template: TemplateSpec, meeting: MeetingModel) -> bytes:
    template_path = Path(template.docx_path)
    if not template_path.is_absolute():
        template_path = (BASE_DIR / template.docx_path).resolve()
    _ensure_template_exists(template_path)

    doc = DocxTemplate(template_path)
    context = _build_context(meeting)
    doc.render(context)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _build_context(meeting: MeetingModel) -> dict:
    contract_dates = _extract_contract_dates(meeting.sections)
    return {
        "project": meeting.meta.project,
        "job_min_no": meeting.meta.job_min_no,
        "description": meeting.meta.description,
        "date": meeting.meta.date,
        "time": meeting.meta.time,
        "location": meeting.meta.location,
        "present": meeting.attendees,
        "apologies": meeting.apologies,
        "distribution": "As above",
        "contract_commencement": contract_dates.contract_commencement,
        "section1_completion": contract_dates.section1_completion,
        "section2_completion": contract_dates.section2_completion,
        "section3_completion": contract_dates.section3_completion,
        "practical_completion": contract_dates.practical_completion,
        "items": _build_items(meeting.sections),
    }


def _build_items(sections: list[Section]) -> list[dict]:
    items: list[dict] = []
    for section in sections:
        first_action = section.actions[0] if section.actions else None
        body = section.title
        if section.notes:
            body = f"{body}\n\n{section.notes}"

        action_summary = ""
        action_due = ""
        if first_action:
            if first_action.owner:
                action_summary = f"{first_action.owner} – {first_action.action}"
            else:
                action_summary = first_action.action
            action_due = first_action.due_date or ""

        items.append(
            {
                "code": section.code,
                "body": body,
                "action_summary": action_summary,
                "action_due": action_due,
            }
        )
    return items


def _extract_contract_dates(sections: list[Section]) -> SectionDates:
    fallback: SectionDates | None = None
    for section in sections:
        if not section.dates:
            continue
        if "contract" in section.title.lower():
            return section.dates
        if fallback is None:
            fallback = section.dates
    return fallback or SectionDates()


def _ensure_template_exists(path: Path) -> None:
    """Create a lightweight fallback template if none is provided."""
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        return

    doc = Document()
    doc.add_heading("Construction Progress Meeting Minutes", level=1)

    meta_labels = [
        ("Project", "{{ project }}"),
        ("Job Min No", "{{ job_min_no }}"),
        ("Description", "{{ description }}"),
        ("Date", "{{ date }}"),
        ("Time", "{{ time }}"),
        ("Location", "{{ location }}"),
    ]
    table = doc.add_table(rows=len(meta_labels), cols=2)
    for row, (label, placeholder) in zip(table.rows, meta_labels, strict=True):
        row.cells[0].text = label
        row.cells[1].text = placeholder

    doc.add_heading("Attendees", level=2)
    doc.add_paragraph(
        "{% for person in present %}- {{ person.name }}"
        "{% if person.company %} ({{ person.company }}){% endif %}"
        "{% if not loop.last %}\n{% endif %}{% endfor %}"
    )

    doc.add_heading("Apologies", level=2)
    doc.add_paragraph(
        "{% for person in apologies %}- {{ person.name }}"
        "{% if person.company %} ({{ person.company }}){% endif %}"
        "{% if not loop.last %}\n{% endif %}{% endfor %}"
    )

    doc.add_heading("Contract Dates", level=2)
    doc.add_paragraph(
        "Commencement: {{ contract_commencement }}\n"
        "Section 1: {{ section1_completion }}\n"
        "Section 2: {{ section2_completion }}\n"
        "Section 3: {{ section3_completion }}\n"
        "Practical Completion: {{ practical_completion }}"
    )

    doc.add_heading("Meeting Sections", level=2)
    doc.add_paragraph(
        "{% for item in items %}Section {{ item.code }}: {{ item.body }}"
        "{% if item.action_summary %}\nAction: {{ item.action_summary }}{% endif %}"
        "{% if item.action_due %} (Due: {{ item.action_due }}){% endif %}"
        "{% if not loop.last %}\n\n{% endif %}{% endfor %}"
    )

    doc.add_paragraph("Distribution: {{ distribution }}")
    doc.save(path)
